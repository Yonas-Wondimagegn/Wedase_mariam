import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
import os
import speech_recognition as sr
from PIL import Image, ImageTk
import pygame
from docx import Document

class Wedase_Mariam_App:
    def __init__(self, master, docx_path, voice_folder):
        self.master = master
        self.master.title("Wedase Mariam")

        # Load Word document
        try:
            self.doc = Document(docx_path)
        except FileNotFoundError:
            messagebox.showerror("Error", "Word file not found.")
            self.master.destroy()
            return

        self.voice_folder = voice_folder
        self.playing = False
        self.paused = False

        # Canvas for displaying Word document content
        self.canvas = tk.Canvas(self.master, width=800, height=1000)
        self.canvas.grid(row=1, column=0, columnspan=4, sticky="nsew")

        # Voice button
        self.voice_button = ttk.Button(self.master, text="Play", command=self.toggle_voice)
        self.voice_button.grid(row=0, column=0, sticky="ew")

        # Text button
        self.text_button = ttk.Button(self.master, text="Text", command=self.search_text)
        self.text_button.grid(row=0, column=1, sticky="ew")

        # Search button
        self.search_button = ttk.Button(self.master, text="Search", command=self.search_text_voice)
        self.search_button.grid(row=0, column=2, sticky="ew")

        # Restart button
        self.restart_button = ttk.Button(self.master, text="Restart", command=self.restart_voice)
        self.restart_button.grid(row=0, column=3, sticky="ew")

        # First display
        self.current_page = 0
        self.display_page()

    def display_page(self):
        self.canvas.delete("all")

        if self.current_page < len(self.doc.paragraphs):
            paragraph = self.doc.paragraphs[self.current_page]
            text = paragraph.text
            self.canvas.create_text(400, 500, text=text, font=("Helvetica", 12), anchor="center")
        else:
            self.canvas.create_text(400, 500, text="End of document", font=("Helvetica", 12), anchor="center")


    def prev_page(self, event):
        if self.current_page > 0:
            self.current_page -= 1
            self.display_page()

    def next_page(self, event):
        if self.current_page < len(self.doc.paragraphs) - 1:
            self.current_page += 1
            self.display_page()

    def toggle_voice(self):
        if not self.playing:
            self.play_voice()
            self.voice_button.config(text="Pause")
        else:
            if self.paused:
                pygame.mixer.music.unpause()
                self.paused = False
                self.voice_button.config(text="Pause")
            else:
                pygame.mixer.music.pause()
                self.paused = True
                self.voice_button.config(text="Resume")

    def restart_voice(self):
        pygame.mixer.music.stop()
        self.play_voice()

    def play_voice(self):
        voice_file = os.path.join(self.voice_folder, f"page_{self.current_page + 1}.mp3")
        if os.path.exists(voice_file):
            try:
                pygame.mixer.init()
                pygame.mixer.music.load(voice_file)
                pygame.mixer.music.play()
                self.playing = True
            except pygame.error as e:
                print("Error playing voice file:", e)
        else:
            print("Voice file not found:", voice_file)

    def search_text(self):
        text_to_search = simpledialog.askstring("Search Text", "Enter the text to search:")
        if text_to_search is None:
            return

        found = False
        for idx, paragraph in enumerate(self.doc.paragraphs):
            if text_to_search in paragraph.text:
                self.current_page = idx
                self.display_page()
                found = True
                break

        if not found:
            print("Text not found in the document.")

    def search_text_voice(self):
        recognizer = sr.Recognizer()

        with sr.Microphone() as source:
            print("Speak the text you want to search for:")
            audio = recognizer.listen(source)

        try:
            search_text = recognizer.recognize_google(audio, language="en-US")

            found = False
            for idx, paragraph in enumerate(self.doc.paragraphs):
                if search_text in paragraph.text:
                    self.current_page = idx
                    self.display_page()
                    found = True
                    break

            if not found:
                print(f"Text '{search_text}' not found in the document.")
        except sr.UnknownValueError:
            print("Could not understand audio.")
        except ValueError:
            print("Could not convert audio to text.")

def main():
    root = tk.Tk()
    app = Wedase_Mariam_App(root, "wedasemariam.docx", "voices")
    root.mainloop()

if __name__ == "__main__":
    main()

print("Thanks for using our bot")
